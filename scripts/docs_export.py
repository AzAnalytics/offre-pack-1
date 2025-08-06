import os
import zipfile
import pandas as pd
from pathlib import Path
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Cm

# Note: We use style names (not style_id) to avoid deprecated lookups

def export_excel(kpis: dict, dataframes: dict[str, pd.DataFrame], out_xlsx: str) -> None:
    """
    Exporte un classeur Excel résumé incluant plusieurs DataFrames.
    - kpis : dict de KPIs pour un onglet 'KPIs'
    - dataframes : dict d'onglet_name -> DataFrame
    """
    os.makedirs(os.path.dirname(out_xlsx) or '.', exist_ok=True)
    with pd.ExcelWriter(out_xlsx, engine='openpyxl') as writer:
        # Onglet KPI
        pd.DataFrame(
            list(kpis.items()), columns=['KPI', 'Value']
        ).to_excel(writer, sheet_name='KPIs', index=False)
        # Autres onglets
        for sheet, df in dataframes.items():
            df.to_excel(writer, sheet_name=sheet, index=False)


def fill_docx(template_path: str, context: dict, out_docx: str) -> None:
    """
    Rend un document Word via un modèle .docx.
    - template_path : chemin vers votre report_template.docx
    - context       : dict des variables à injecter
    - out_docx      : chemin de sortie du .docx généré
    """
    # Préparation du dossier de sortie
    parent = os.path.dirname(out_docx)
    if parent:
        os.makedirs(parent, exist_ok=True)

    # Racine du projet pour templates locaux
    project_root = Path(__file__).resolve().parent.parent

    # Sélection du template
    tpl_path = Path(template_path)
    if not (tpl_path.is_file() and zipfile.is_zipfile(tpl_path)):
        local_tpl = project_root / 'templates' / 'report_template.docx'
        default_tpl = project_root / 'templates' / 'default_report_template.docx'
        if local_tpl.is_file() and zipfile.is_zipfile(local_tpl):
            tpl_path = local_tpl
            print(f"WARNING: template utilisateur introuvable, utilisation de '{local_tpl}'")
        elif default_tpl.is_file() and zipfile.is_zipfile(default_tpl):
            tpl_path = default_tpl
            print(f"WARNING: aucun template valide trouvé, utilisation de '{default_tpl}'")
        else:
            raise FileNotFoundError(
                f"Aucun template DOCX valide trouvé : '{template_path}', '{local_tpl}', ni '{default_tpl}'"
            )

    # Charger le document
    doc = DocxTemplate(str(tpl_path))

    # Embedding images
    for key, val in list(context.items()):
        if isinstance(val, str) and os.path.splitext(val)[1].lower() in ('.png', '.jpg', '.jpeg'):
            context[key] = InlineImage(doc, val, width=Cm(12))
        elif isinstance(val, (list, tuple)):
            new_list = []
            for item in val:
                if isinstance(item, str) and os.path.splitext(item)[1].lower() in ('.png', '.jpg', '.jpeg'):
                    new_list.append(InlineImage(doc, item, width=Cm(12)))
                else:
                    new_list.append(item)
            context[key] = new_list

    # Render Jinja context
    doc.render(context)

        # Insérer les DataFrames sous forme de tableaux Word
    dataframes = context.get('dataframes', {})
    for sheet_name, df in dataframes.items():
        # Titre de section : tentative avec style 'Heading 2', fallback au style par défaut
        try:
            paragraph = doc.add_paragraph(sheet_name, style='Heading 2')
        except (KeyError, ValueError):
            paragraph = doc.add_paragraph(sheet_name)
        # Création du tableau
        table = doc.add_table(rows=1, cols=len(df.columns))
        # Applique un style de tableau si disponible
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
        # En-têtes
        hdr_cells = table.rows[0].cells
        for idx, col in enumerate(df.columns):
            hdr_cells[idx].text = str(col)
        # Lignes de données
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for j, val in enumerate(row):
                row_cells[j].text = str(val)

    # Sauvegarder le document
    doc.save(out_docx)
